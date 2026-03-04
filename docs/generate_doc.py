"""Generate the Torah Chat Technical Reference Word document."""

from docx import Document
from docx.shared import Inches, Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import os

doc = Document()

# ── Styles ─────────────────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

# Code style
if 'Code' not in [s.name for s in doc.styles]:
    code_style = doc.styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
    code_style.font.name = 'Consolas'
    code_style.font.size = Pt(9)
    code_style.paragraph_format.space_before = Pt(2)
    code_style.paragraph_format.space_after = Pt(2)

def add_code_block(text):
    """Add a formatted code block."""
    for line in text.strip().split('\n'):
        p = doc.add_paragraph(line, style='Code')
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

def add_table(headers, rows):
    """Add a formatted table."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        hdr.cells[i].text = h
        for p in hdr.cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = str(val)
            for p in row.cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    return table

def add_image_safe(path, width=Inches(6)):
    """Add image if it exists."""
    if os.path.exists(path):
        doc.add_picture(path, width=width)
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ══════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════

doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Torah Chat')
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Technical Reference & Architecture Guide')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()

version = doc.add_paragraph()
version.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = version.add_run('Version 1.3.2')
run.font.size = Pt(14)

date = doc.add_paragraph()
date.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date.add_run('February 2026')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_paragraph()
disclaimer = doc.add_paragraph()
disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = disclaimer.add_run('© 2026 Jason Leznek · MIT License')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

note = doc.add_paragraph()
note.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = note.add_run('Not developed by or affiliated with Sefaria.org')
run.font.size = Pt(9)
run.font.italic = True
run.font.color.rgb = RGBColor(0xaa, 0xaa, 0xaa)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ══════════════════════════════════════════════════════════════════════
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Project Overview',
    '2. Architecture',
    '3. Application Flowchart',
    '4. Message Data Flow (Sequence Diagram)',
    '5. IPC API Reference',
    '6. Provider Plugin Guide',
    '7. MCP Tool Catalog',
    '8. Build & Release Guide',
    '9. Configuration Reference',
    '10. Security Model',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(4)
    p.runs[0].font.size = Pt(12)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# 1. PROJECT OVERVIEW
# ══════════════════════════════════════════════════════════════════════
doc.add_heading('1. Project Overview', level=1)

doc.add_paragraph(
    'Torah Chat is a standalone Electron desktop application for exploring the Sefaria digital library '
    'of Jewish texts. It connects to Sefaria\'s MCP (Model Context Protocol) servers and supports '
    'multiple AI providers — Google Gemini, OpenAI, Anthropic Claude, and Ollama (local) — for '
    'AI-powered chat with tool calling.'
)

doc.add_heading('Key Capabilities', level=2)
bullets = [
    'Multi-provider AI chat with streaming responses and live Markdown rendering',
    'Automatic tool calling against Sefaria\'s MCP servers for accurate text retrieval',
    'Auto-linking of Sefaria citations in responses (works with all providers including local models)',
    'Chat history with auto-save and restore',
    'In-app text viewer — click any citation to read the source in a side pane',
    'Print preview with PDF generation',
    'Mermaid diagram and LaTeX math rendering in responses',
    'Auto-update with progress bar via GitHub Releases',
    'Code-signed Windows installers (x64 + arm64)',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('Technology Stack', level=2)
add_table(
    ['Component', 'Technology'],
    [
        ['Runtime', 'Electron (Chromium + Node.js)'],
        ['Language', 'TypeScript (main/preload), Vanilla JS (renderer)'],
        ['AI Providers', 'Google Gemini, OpenAI, Anthropic, Ollama'],
        ['MCP Client', '@modelcontextprotocol/sdk (SSE + Streamable HTTP)'],
        ['Bundler', 'esbuild (main + preload bundles)'],
        ['Packaging', 'electron-builder → Microsoft Store (appx)'],
    ]
)

doc.add_heading('Key Files', level=2)
add_table(
    ['File', 'Purpose'],
    [
        ['src/main.ts', 'Electron main process: window creation, IPC handlers, settings persistence'],
        ['src/preload.ts', 'Context bridge exposing window.sefaria API to the renderer'],
        ['src/chat-engine.ts', 'LLM streaming + tool-calling loop (up to 10 rounds)'],
        ['src/mcp-client.ts', 'Manages SSE connections to both Sefaria MCP servers'],
        ['src/prompts.ts', 'System and command prompts'],
        ['src/providers/', 'Pluggable LLM provider implementations (Gemini, OpenAI, Anthropic, Ollama)'],
        ['src/renderer/', 'Browser-side UI (HTML, CSS, vanilla JS)'],
        ['esbuild.js', 'Build script for main + preload bundles'],
        ['package.json', 'Scripts, deps, electron-builder config'],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# 2. ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════
doc.add_heading('2. Architecture', level=1)

doc.add_paragraph(
    'The application follows Electron\'s multi-process architecture with strict separation '
    'between the main process (Node.js) and the renderer process (Chromium). All communication '
    'passes through a preload script context bridge, ensuring the renderer never has direct '
    'access to Node.js APIs.'
)

add_image_safe('docs/architecture.png', Inches(6.5))

doc.add_heading('Process Model', level=2)

doc.add_heading('Main Process (main.ts)', level=3)
doc.add_paragraph(
    'The main process runs in Node.js and manages the application lifecycle, window creation, '
    'settings persistence, chat history, MCP connections, the chat engine, and auto-updates. '
    'All IPC handlers are registered here.'
)

doc.add_heading('Preload Script (preload.ts)', level=3)
doc.add_paragraph(
    'The preload script runs in a sandboxed context with access to both Node.js and the renderer DOM. '
    'It exposes a typed window.sefaria API via Electron\'s contextBridge, mapping each method to an '
    'IPC invoke or event listener.'
)

doc.add_heading('Renderer Process (renderer/)', level=3)
doc.add_paragraph(
    'The renderer is a single-page application built with vanilla HTML, CSS, and JavaScript — no '
    'framework. It handles the chat UI, settings panel, embedded browser pane, Markdown rendering '
    '(with Mermaid and KaTeX), and auto-linking of Sefaria citations.'
)

doc.add_heading('External Services', level=2)
add_table(
    ['Service', 'URL', 'Purpose'],
    [
        ['Sefaria Texts MCP', 'https://mcp.sefaria.org/sse', 'Query the Sefaria library of Jewish texts'],
        ['Sefaria Developers MCP', 'https://developers.sefaria.org/mcp', 'API/code developer assistance'],
        ['Google Gemini API', 'generativelanguage.googleapis.com', 'Gemini LLM inference'],
        ['OpenAI API', 'api.openai.com', 'GPT/o-series LLM inference'],
        ['Anthropic API', 'api.anthropic.com', 'Claude LLM inference'],
        ['Ollama (local)', 'localhost:11434', 'Local LLM inference (no internet)'],
        ['GitHub Releases', 'github.com/jleznek/torah-chat', 'Auto-update distribution'],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# 3. APPLICATION FLOWCHART
# ══════════════════════════════════════════════════════════════════════
doc.add_heading('3. Application Flowchart', level=1)

doc.add_paragraph(
    'The following diagram shows the three main flows through the application: '
    'the settings configuration flow, the chat message flow (with tool-calling loop), '
    'and the auto-update flow.'
)

add_image_safe('docs/flowchart.png', Inches(6.5))

doc.add_heading('Chat Message Flow Detail', level=2)
steps = [
    ('1. User input', 'User types a message and presses Enter (or clicks Send).'),
    ('2. Rate limit check', 'The ChatEngine checks if the request is within the provider\'s RPM limit. If at capacity, it waits until a slot opens.'),
    ('3. Stream to LLM', 'The message (with full conversation history and system prompt) is sent to the configured LLM provider. Text is streamed back in real-time.'),
    ('4. Tool-calling loop', 'If the LLM response includes function calls, the engine executes each one against the appropriate Sefaria MCP server, adds results to the history, and sends it back to the LLM. This repeats for up to 10 rounds.'),
    ('5. Final render', 'The complete response is rendered as Markdown with automatic Sefaria citation linking, Mermaid diagrams, and KaTeX math.'),
    ('6. Auto-save', 'The conversation is automatically saved to disk as a JSON file.'),
    ('7. Follow-ups', 'If rate limit budget allows, the engine generates 2-3 follow-up question suggestions.'),
]
for title, desc in steps:
    p = doc.add_paragraph()
    run = p.add_run(title + ': ')
    run.bold = True
    p.add_run(desc)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# 4. SEQUENCE DIAGRAM
# ══════════════════════════════════════════════════════════════════════
doc.add_heading('4. Message Data Flow (Sequence Diagram)', level=1)

doc.add_paragraph(
    'This sequence diagram traces a single user message through every layer of the application, '
    'showing IPC communication, LLM streaming, tool-call execution, and response rendering.'
)

add_image_safe('docs/sequence.png', Inches(6.5))

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# 5. IPC API REFERENCE
# ══════════════════════════════════════════════════════════════════════
doc.add_heading('5. IPC API Reference', level=1)

doc.add_paragraph(
    'All communication between the renderer and main process passes through Electron\'s IPC system. '
    'The preload script exposes these as the window.sefaria API. Below is the complete reference.'
)

doc.add_heading('Renderer → Main (invoke/handle)', level=2)
doc.add_paragraph(
    'These channels use ipcRenderer.invoke() / ipcMain.handle() (request-response pattern).'
)

ipc_r2m = [
    ['get-providers', '(none)', 'ProviderInfo[]', 'List all available AI providers'],
    ['get-provider-config', '(none)', '{providerId, modelId, hasKey}', 'Get active provider config'],
    ['save-provider-config', '{providerId, modelId, apiKey?}', 'true', 'Save provider settings, reinit engine'],
    ['get-configured-providers', '(none)', '(ProviderInfo & {hasKey})[]', 'All providers with key status'],
    ['switch-provider', '{providerId, modelId}', '{success?, error?}', 'Quick-switch without clearing history'],
    ['remove-provider-key', 'providerId', '{success?, switchedTo?, modelId?}', 'Remove a provider\'s API key'],
    ['detect-ollama', '(none)', '{available, models[]}', 'Detect local Ollama installation'],
    ['get-api-key', '(none)', 'string', 'Legacy: get active provider key'],
    ['set-api-key', 'apiKey', 'true', 'Legacy: set active provider key'],
    ['get-mcp-status', '(none)', '{connected, toolCount, servers?}', 'MCP connection status'],
    ['send-message', '{message, responseLength?}', '{success?, error?, chatId?}', 'Send message, triggers streaming'],
    ['clear-chat', '(none)', 'true', 'Clear conversation history'],
    ['reconnect-mcp', '(none)', 'true', 'Reconnect MCP servers'],
    ['list-chats', '(none)', 'ChatSummary[]', 'List saved chats (most recent first)'],
    ['load-chat', 'chatId', 'SavedChat | null', 'Load and restore a saved chat'],
    ['delete-chat', 'chatId', 'true', 'Delete a saved chat'],
    ['new-chat', '(none)', 'true', 'Start fresh chat'],
    ['get-usage-stats', '(none)', '{used, limit, resetsInSeconds}', 'Rate limit usage'],
    ['resize-for-webview', 'open: boolean', 'void', 'Resize window for side pane'],
    ['print-chat', '{html}', 'void', 'Generate PDF from chat HTML'],
    ['get-app-version', '(none)', 'string', 'App version from package.json'],
    ['get-changelog', '(none)', 'string', 'CHANGELOG.md contents'],
]
add_table(['Channel', 'Parameters', 'Returns', 'Description'], ipc_r2m)

doc.add_paragraph()
doc.add_heading('Main → Renderer (events)', level=2)
doc.add_paragraph(
    'These channels use webContents.send() / ipcRenderer.on() (push events from main to renderer).'
)

ipc_m2r = [
    ['chat-stream', '{chunk: string}', 'Incremental text chunk during streaming'],
    ['chat-stream-end', '{followUps?: string[]}', 'End of stream; includes follow-up suggestions'],
    ['tool-status', '{toolName, status}', 'MCP tool call progress (calling/done)'],
    ['mcp-status', '{connected, toolCount?, error?, servers?}', 'MCP connection status update'],
    ['open-url', 'url: string', 'Open URL in embedded webview pane'],
    ['usage-update', '{used, limit, resetsInSeconds}', 'Rate limit stats after each message'],
]
add_table(['Channel', 'Data', 'Description'], ipc_m2r)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# 6. PROVIDER PLUGIN GUIDE
# ══════════════════════════════════════════════════════════════════════
doc.add_heading('6. Provider Plugin Guide', level=1)

doc.add_paragraph(
    'Adding a new AI provider requires implementing the ChatProvider interface and registering it '
    'in the provider index. The architecture is designed to make this straightforward.'
)

doc.add_heading('ChatProvider Interface', level=2)
add_code_block('''interface ChatProvider {
    readonly info: ProviderInfo;
    
    streamChat(
        history: Message[],
        systemPrompt: string,
        tools: ToolDeclaration[],
        onTextChunk: (text: string) => void
    ): Promise<StreamResult>;
    
    generateOnce(prompt: string): Promise<string>;
}''')

doc.add_heading('ProviderInfo Interface', level=2)
add_code_block('''interface ProviderInfo {
    id: string;              // e.g. 'gemini', 'openai'
    name: string;            // Display name
    models: ProviderModel[]; // Available models
    defaultModel: string;    // Default model ID
    rateLimit: { rpm: number; windowMs: number };
    requiresKey?: boolean;   // false for Ollama
    keyPlaceholder: string;  // e.g. 'AIza...'
    keyHelpUrl: string;      // Link to get an API key
    keyHelpLabel: string;    // e.g. 'Google AI Studio'
}''')

doc.add_heading('Steps to Add a New Provider', level=2)
steps = [
    'Create src/providers/myprovider.ts implementing the ChatProvider interface',
    'Export a MY_PROVIDER_INFO constant with the provider metadata (ProviderInfo)',
    'Export a MyProvider class with streamChat() and generateOnce() methods',
    'Add the provider to src/providers/index.ts: import it, add to AVAILABLE_PROVIDERS array, add a case to createProvider()',
    'Test: npm start → Settings → select your new provider → enter API key → send a message',
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {step}')

doc.add_heading('Existing Providers', level=2)
add_table(
    ['Provider', 'ID', 'Default Model', 'RPM Limit', 'Key Required'],
    [
        ['Google Gemini', 'gemini', 'gemini-2.5-flash', '5 (model-dependent)', 'Yes (free tier)'],
        ['OpenAI', 'openai', 'gpt-4.1-mini', '500', 'Yes (paid)'],
        ['Anthropic', 'anthropic', 'claude-sonnet-4-20250514', '50', 'Yes (paid)'],
        ['Ollama', 'ollama', 'llama3.2', '9999 (local)', 'No'],
    ]
)

doc.add_heading('Key Implementation Notes', level=2)
notes = [
    'streamChat() must handle both text streaming AND tool-call extraction from the LLM response',
    'Tool declarations must be converted from the MCP format to the provider\'s native format',
    'Function call results must be formatted back into the provider\'s expected message format',
    'generateOnce() is used for follow-up question generation (no tool calling needed)',
    'Rate limiting is handled by the ChatEngine, not the provider',
]
for n in notes:
    doc.add_paragraph(n, style='List Bullet')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# 7. MCP TOOL CATALOG
# ══════════════════════════════════════════════════════════════════════
doc.add_heading('7. MCP Tool Catalog', level=1)

doc.add_paragraph(
    'The app connects to two Sefaria MCP servers at startup. Tools are discovered dynamically '
    'via the MCP protocol. Below are the known tools as of this writing.'
)

doc.add_heading('Sefaria Texts MCP (mcp.sefaria.org)', level=2)
doc.add_paragraph(
    'This server provides tools for querying the Sefaria library of Jewish texts.'
)
texts_tools = [
    ['get_text', 'Retrieve the text of a specific Sefaria reference', 'ref: string (e.g. "Genesis 1:1")'],
    ['english_semantic_search', 'Semantic search across English translations', 'query: string, filters?: object'],
    ['text_search', 'Full-text search across the library', 'query: string, path?: string'],
    ['search_in_book', 'Search within a specific book', 'book: string, query: string'],
    ['get_links_between_texts', 'Get links/connections between two texts', 'ref1: string, ref2: string'],
    ['get_english_translations', 'Get available English translations', 'ref: string'],
    ['get_text_catalogue_info', 'Get metadata about a text', 'ref: string'],
    ['get_text_or_category_shape', 'Get structure/shape of a text or category', 'ref: string'],
    ['get_current_calendar', 'Get today\'s calendar readings', '(none)'],
    ['get_topic_details', 'Get details about a topic', 'topic: string'],
    ['search_in_dictionaries', 'Search dictionary entries', 'query: string'],
    ['get_available_manuscripts', 'List available manuscripts', 'ref: string'],
    ['get_manuscript_image', 'Get manuscript image URL', 'manuscript_id: string'],
    ['clarify_name_argument', 'Disambiguate a text reference', 'name: string'],
    ['clarify_search_path_filter', 'Help build a search filter path', 'query: string'],
]
add_table(['Tool', 'Description', 'Key Parameters'], texts_tools)

doc.add_paragraph()
doc.add_heading('Sefaria Developers MCP (developers.sefaria.org)', level=2)
doc.add_paragraph(
    'This server provides tools for querying the Sefaria API documentation, '
    'helping developers build applications with the Sefaria API.'
)

doc.add_heading('How Tool Calling Works', level=2)
steps = [
    'On startup, McpClientManager connects to both servers and discovers available tools via the MCP listTools protocol',
    'When the ChatEngine sends a request to the LLM, it includes all discovered tool declarations',
    'The LLM can request tool calls in its response (e.g. get_text with ref "Genesis 1:1")',
    'The ChatEngine routes each call to the correct MCP server via McpClientManager.callTool()',
    'Tool results are added to the conversation history and sent back to the LLM',
    'This loop continues for up to 10 rounds until the LLM produces a final text response',
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {s}')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# 8. BUILD & RELEASE GUIDE
# ══════════════════════════════════════════════════════════════════════
doc.add_heading('8. Build & Release Guide', level=1)

doc.add_heading('Development Setup', level=2)
add_code_block('''npm install          # Install dependencies
npm start            # Build and launch
npm run watch        # Watch mode (auto-rebuild on changes)
npx electron .       # Launch without rebuilding
# Press F5 in VS Code to debug the main process''')

doc.add_heading('Build Pipeline', level=2)
doc.add_paragraph('The build process has three stages:')
steps = [
    ('Type check', 'tsc --noEmit — validates TypeScript without emitting'),
    ('Bundle', 'esbuild bundles src/main.ts → dist/main.js and src/preload.ts → dist/preload.js'),
    ('Copy renderer', 'index.html, styles.css, and renderer.js are copied to dist/renderer/'),
]
for title, desc in steps:
    p = doc.add_paragraph()
    run = p.add_run(f'{title}: ')
    run.bold = True
    p.add_run(desc)

doc.add_heading('Building Distributables', level=2)
add_code_block('''# Windows (NSIS installer + portable, x64 + arm64)
npm run dist:win

# macOS .dmg
npm run dist:mac

# Linux AppImage + .deb
npm run dist:linux''')

doc.add_paragraph('Output goes to the release/ directory.')

doc.add_heading('Release Artifacts', level=2)
add_table(
    ['Artifact', 'Type', 'Architectures'],
    [
        ['Torah Chat-Setup-{ver}.exe', 'NSIS installer (universal)', 'x64 + arm64'],
        ['Torah Chat-Setup-{ver}-x64.exe', 'NSIS installer', 'x64 only'],
        ['Torah Chat-Setup-{ver}-arm64.exe', 'NSIS installer', 'arm64 only'],
        ['Torah Chat-Portable-{ver}.exe', 'Portable (universal)', 'x64 + arm64'],
        ['Torah Chat-Portable-{ver}-x64.exe', 'Portable', 'x64 only'],
        ['Torah Chat-Portable-{ver}-arm64.exe', 'Portable', 'arm64 only'],
        ['latest.yml', 'Auto-updater manifest', 'N/A'],
        ['*.blockmap', 'Differential update data', 'Per installer'],
    ]
)

doc.add_heading('Publishing a Release', level=2)
steps = [
    'Update version in package.json',
    'Add entry to CHANGELOG.md',
    'Commit and push: git add -A && git commit -m "v{X.Y.Z}: description" && git push',
    'Tag: git tag v{X.Y.Z} && git push --tags',
    'Build: npm run dist:win',
    'Create GitHub Release with all artifacts:',
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {s}')

add_code_block('''gh release create v{X.Y.Z} --title "v{X.Y.Z}" --notes "release notes" \\
  release/Sefaria\\ Chat-Setup-*.exe \\
  release/Sefaria\\ Chat-Portable-*.exe \\
  release/latest.yml \\
  release/*.blockmap''')

doc.add_heading('Code Signing', level=2)
doc.add_paragraph(
    'Windows builds are automatically signed using a certificate configured in package.json '
    '(signtoolOptions). The certificate must be installed in the Windows Certificate Store '
    '(Current User/Personal) with subject name "Torah Chat".'
)

doc.add_heading('Distribution & Updates', level=2)
doc.add_paragraph(
    'The app is distributed exclusively through the Microsoft Store. Updates are delivered '
    'automatically by the Store. The appx package is built using electron-builder.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# 9. CONFIGURATION REFERENCE
# ══════════════════════════════════════════════════════════════════════
doc.add_heading('9. Configuration Reference', level=1)

doc.add_heading('Settings File', level=2)
doc.add_paragraph('Location: {userData}/settings.json (typically %APPDATA%/sefaria on Windows)')
add_code_block('''{
    "provider": "gemini",                    // Active provider ID
    "model": "gemini-2.5-flash",            // Active model ID
    "apiKeys": {                             // API keys by provider
        "gemini": "AIza...",
        "openai": "sk-...",
        "anthropic": "sk-ant-..."
    },
    "windowBounds": {                        // Window position/size
        "x": 100, "y": 100,
        "width": 1200, "height": 800
    },
    "windowMaximized": false
}''')

doc.add_heading('Chat History', level=2)
doc.add_paragraph('Location: {userData}/chats/')
doc.add_paragraph('Each chat is saved as a JSON file with the following structure:')
add_code_block('''{
    "id": "chat_1707849600000_abc123",       // Unique ID
    "title": "What is the Shema?",           // First user message (≤60 chars)
    "createdAt": "2026-02-13T10:00:00.000Z",
    "updatedAt": "2026-02-13T10:05:00.000Z",
    "messages": [...],                        // UI-visible messages
    "history": [...]                          // Full engine history (for resuming)
}''')

doc.add_heading('Chat ID Format', level=2)
doc.add_paragraph(
    'Chat IDs follow the pattern chat_{timestamp}_{random6chars} and are validated with the regex '
    '/^chat_\\d+_[a-z0-9]+$/ to prevent path traversal attacks.'
)

doc.add_heading('User Data Directory', level=2)
add_table(
    ['Path', 'Contents'],
    [
        ['{userData}/settings.json', 'Provider config, API keys, window state'],
        ['{userData}/chats/', 'Saved chat JSON files'],
        ['{userData}/logs/', 'Electron log files (if enabled)'],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# 10. SECURITY MODEL
# ══════════════════════════════════════════════════════════════════════
doc.add_heading('10. Security Model', level=1)

doc.add_paragraph(
    'The application follows Electron security best practices to minimize attack surface. '
    'Below is a summary of the security controls in place.'
)

doc.add_heading('Electron Security Configuration', level=2)
add_table(
    ['Setting', 'Value', 'Purpose'],
    [
        ['contextIsolation', 'true', 'Renderer code cannot access Node.js globals'],
        ['nodeIntegration', 'false', 'No require() or process in renderer'],
        ['sandbox', 'true', 'Renderer runs in Chromium sandbox'],
        ['webviewTag', 'true', 'Needed for embedded text viewer (restricted)'],
    ]
)

doc.add_heading('Context Bridge', level=2)
doc.add_paragraph(
    'The preload script exposes a strictly typed window.sefaria API via contextBridge. '
    'The renderer has no access to ipcRenderer, require, process, or any Node.js API directly. '
    'Each API method maps to a specific IPC channel with defined parameter types.'
)

doc.add_heading('Navigation & Popup Control', level=2)
controls = [
    'setWindowOpenHandler: Denies all popup windows; external URLs are routed to the embedded webview',
    'will-navigate: Blocks navigation away from the app\'s file:// URL; external URLs open in webview',
    'Link sanitization: The renderer strips non-http(s)/mailto hrefs, preventing javascript: and data: URI attacks',
    'Menu.setApplicationMenu(null): Removes default Electron menu to prevent unintended actions',
]
for c in controls:
    doc.add_paragraph(c, style='List Bullet')

doc.add_heading('Input Validation', level=2)
validations = [
    'Chat ID validation: isValidChatId() uses regex /^chat_\\d+_[a-z0-9]+$/ to prevent path traversal when loading/deleting chats',
    'API keys: Stored locally in {userData}/settings.json; never transmitted except to the configured LLM API endpoint',
    'MCP server URLs: Hardcoded to https://mcp.sefaria.org/sse and https://developers.sefaria.org/mcp — not user-configurable',
    'Rate limiting: Client-side RPM tracking prevents API abuse and protects against runaway tool-call loops (max 10 rounds)',
]
for v in validations:
    doc.add_paragraph(v, style='List Bullet')

doc.add_heading('Data Privacy', level=2)
doc.add_paragraph(
    'All user data (settings, API keys, chat history) is stored locally on the user\'s machine '
    'in the Electron userData directory. No telemetry or analytics are collected. Chat messages are '
    'sent only to the user-configured LLM provider and Sefaria\'s MCP servers.'
)

doc.add_heading('Code Signing', level=2)
doc.add_paragraph(
    'Windows distributables are signed with a SHA-256 certificate (subject: "Torah Chat") '
    'to ensure binary integrity and prevent tampering warnings. The certificate is stored in '
    'the Windows Certificate Store, not in the repository.'
)

doc.add_heading('Update Security', level=2)
doc.add_paragraph(
    'Updates are delivered through the Microsoft Store, which handles code signing verification '
    'and integrity checks. All communication uses HTTPS.'
)

# ── Save ───────────────────────────────────────────────────────────────
output_path = os.path.join('docs', 'Torah Chat - Technical Reference.docx')
doc.save(output_path)
print(f'Document saved to: {output_path}')
