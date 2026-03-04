"""Generate a formatted Word document announcement for Torah Chat."""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# -- Page margins --
for section in doc.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

# -- Style tweaks --
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(6)

# Heading styles
for level, size, color in [
    ('Heading 1', 22, RGBColor(0x1A, 0x47, 0x8A)),
    ('Heading 2', 15, RGBColor(0x2C, 0x5F, 0xA1)),
    ('Heading 3', 13, RGBColor(0x3A, 0x7C, 0xBD)),
]:
    h = doc.styles[level]
    h.font.name = 'Calibri'
    h.font.size = Pt(size)
    h.font.color.rgb = color
    h.font.bold = True
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after = Pt(6)


def add_styled_paragraph(text, bold=False, italic=False, size=None, color=None, alignment=None, space_after=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if alignment is not None:
        p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p


def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def set_cell_shading(cell, color_hex):
    """Set background shading on a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear',
    })
    shading.append(shading_elem)


# ============================================================
# TITLE
# ============================================================
title = doc.add_heading('Torah Chat', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1A, 0x47, 0x8A)

subtitle = add_styled_paragraph(
    'AI-Powered Jewish Text Exploration, Right on Your Desktop',
    bold=True, size=14, color=RGBColor(0x55, 0x55, 0x55),
    alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4
)

# Decorative line
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('━' * 50)
run.font.color.rgb = RGBColor(0x1A, 0x47, 0x8A)
run.font.size = Pt(10)
p.paragraph_format.space_after = Pt(12)

# ============================================================
# INTRO
# ============================================================
p = doc.add_paragraph()
run = p.add_run('Ever wished you could have a conversation with the entire Sefaria library? ')
run.italic = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
run = p.add_run('Now you can.')
run.bold = True
run.italic = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x1A, 0x47, 0x8A)

doc.add_paragraph(
    'Torah Chat is a free, open-source desktop app that connects the world\u2019s largest '
    'digital collection of Jewish texts \u2014 Torah, Talmud, Midrash, Halakha, Kabbalah, '
    'philosophy, and more \u2014 to the AI model of your choice. Ask a question in plain English, '
    'and the AI automatically searches Sefaria\u2019s library, retrieves primary sources with '
    'original Hebrew/Aramaic alongside English translations, and weaves them into rich, scholarly '
    'responses \u2014 complete with clickable links back to Sefaria.org.'
)

# ============================================================
# WHAT CAN YOU ASK?
# ============================================================
doc.add_heading('What Can You Ask?', level=2)

prompts = [
    ('\u201cWhat is today\u2019s Torah portion and Daf Yomi?\u201d',
     ' \u2014 pulls the live Jewish calendar'),
    ('\u201cShow me an ancient manuscript of Genesis 1:1\u201d',
     ' \u2014 retrieves and displays actual manuscript images'),
    ('\u201cCompare translations of Psalm 23\u201d',
     ' \u2014 shows multiple English renderings side by side'),
    ('\u201cWhat does Judaism teach about forgiveness and repentance?\u201d',
     ' \u2014 searches across the entire library and synthesizes sources from Tanakh, Talmud, Rambam, and beyond'),
    ('\u201cExplain the structure of the Talmud \u2014 Mishnah and Gemara\u201d',
     ' \u2014 gives a comprehensive educational overview'),
    ('\u201cTell me about Shabbat \u2014 its sources, themes, and subtopics\u201d',
     ' \u2014 pulls Sefaria\u2019s curated topic pages'),
    ('\u201cLook up the word \u2018chesed\u2019 in Jastrow\u2019s dictionary\u201d',
     ' \u2014 searches classical Hebrew/Aramaic dictionaries'),
    ('\u201cWhat does Rashi say about Noah\u2019s drunkenness?\u201d',
     ' \u2014 pinpoints specific commentary'),
]

for prompt_text, description in prompts:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(prompt_text)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x47, 0x8A)
    p.add_run(description)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Every text reference is a clickable hyperlink. ')
run.bold = True
p.add_run('Click any source and it opens right inside the app in a side-by-side embedded browser.')

# ============================================================
# USE ANY AI MODEL
# ============================================================
doc.add_heading('Use Any AI Model \u2014 Online or Offline', level=2)

doc.add_paragraph(
    'Torah Chat supports 9 LLM providers with dozens of models, so you pick what works for you:'
)

# Provider table
table = doc.add_table(rows=10, cols=3)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['', 'Provider', 'Highlights']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(10)

providers = [
    ('\u2601\ufe0f', 'Google Gemini', 'Gemini 2.5 Flash, Pro, and more \u2014 free tier available!'),
    ('\u2601\ufe0f', 'OpenAI', 'GPT-4.1, GPT-4o, o4 Mini'),
    ('\u2601\ufe0f', 'Anthropic', 'Claude Sonnet 4, Claude 3.5 Haiku'),
    ('\u2601\ufe0f', 'xAI', 'Grok 3, Grok 3 Mini (fast & regular)'),
    ('\u2601\ufe0f', 'Mistral AI', 'Mistral Small, Medium, Large'),
    ('\u2601\ufe0f', 'DeepSeek', 'DeepSeek-V3, DeepSeek-R1'),
    ('\u2601\ufe0f', 'Groq', 'Llama 3.3 70B, Mixtral, Gemma 2 \u2014 blazing fast inference'),
    ('\u2601\ufe0f', 'OpenRouter', 'Multi-model gateway \u2014 access dozens of models through one key'),
    ('\U0001f4bb', 'Ollama (Local)', 'Runs 100% offline. No API key. No internet. Auto-detects your models.'),
]

for row_idx, (icon, provider, highlights) in enumerate(providers, start=1):
    table.rows[row_idx].cells[0].text = icon
    table.rows[row_idx].cells[1].text = provider
    table.rows[row_idx].cells[2].text = highlights
    for cell_idx in range(3):
        for paragraph in table.rows[row_idx].cells[cell_idx].paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)
    # Bold provider name
    for paragraph in table.rows[row_idx].cells[1].paragraphs:
        for run in paragraph.runs:
            run.bold = True

# Set column widths
for row in table.rows:
    row.cells[0].width = Inches(0.4)
    row.cells[1].width = Inches(1.6)
    row.cells[2].width = Inches(4.5)

doc.add_paragraph()

# Callout paragraphs
p = doc.add_paragraph()
run = p.add_run('Want to use a free model? ')
run.bold = True
p.add_run('Google Gemini\u2019s free tier gives you powerful AI at zero cost. ')
run = p.add_run('Want total privacy? ')
run.bold = True
p.add_run('Ollama runs entirely on your machine \u2014 your questions never leave your computer.')

# ============================================================
# KEY FEATURES
# ============================================================
doc.add_heading('Key Features', level=2)

features = [
    ('Streaming responses', ' \u2014 watch the AI think in real time with live Markdown rendering'),
    ('Automatic source retrieval', ' \u2014 the AI calls up to 10 rounds of tools to find exactly the right texts'),
    ('Mermaid diagrams', ' \u2014 ask for a timeline or flowchart and get a visual diagram'),
    ('Citation auto-linking', ' \u2014 even bare references like \u201cBerakhot 2a\u201d become clickable Sefaria links'),
    ('Chat history', ' \u2014 save, load, and manage multiple conversations'),
    ('Print to PDF', ' \u2014 export any conversation for study or sharing'),
    ('First-run setup wizard', ' \u2014 choose your provider and start chatting in under a minute'),
    ('Auto-updates', ' \u2014 always get the latest features automatically'),
    ('100% private', ' \u2014 no telemetry, no analytics, no tracking. Your API keys and chats stay on your machine.'),
]

for feat_name, feat_desc in features:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(feat_name)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x47, 0x8A)
    p.add_run(feat_desc)

# ============================================================
# GET STARTED
# ============================================================
doc.add_heading('Get Started', level=2)

steps = [
    'Download Torah Chat (available on Windows, including the Microsoft Store)',
    'Pick a provider (try Gemini for free, or Ollama for fully offline)',
    'Start asking questions!',
]
for i, step in enumerate(steps, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'{i}.  ')
    run.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x47, 0x8A)
    run.font.size = Pt(12)
    run2 = p.add_run(step)
    run2.font.size = Pt(12)

doc.add_paragraph()

# Closing
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Built with love for the Jewish text tradition.')
run.italic = True
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('MIT licensed \u2022 Open source \u2022 Free forever')
run.bold = True
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x1A, 0x47, 0x8A)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Torah Chat is an independent project and is not developed by or affiliated with Sefaria.org.')
run.italic = True
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# Save
output_path = os.path.join(os.path.dirname(__file__), 'Torah Chat Announcement.docx')
doc.save(output_path)
print(f'Saved: {output_path}')
